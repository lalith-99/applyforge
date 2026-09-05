"""PDF/DOCX resume rendering from a structured ResumeProfile (see
MASTER_REQUIREMENTS.md §36-§38). Deterministic, template-based rendering —
no AI involved once the profile content itself has been produced upstream.
"""

from __future__ import annotations

import io
import unicodedata

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, RGBColor
from fpdf import FPDF

from app.resume.models import ContactInfo, ExperienceEntry, ResumeProfile

# fpdf2's built-in core fonts default to latin-1, but cp1252 (a superset)
# natively supports the bullet/dash/smart-quote characters this template
# uses, so we only need NFKD decomposition for ligatures (e.g. the
# single-glyph "fi" in "Certifications") rather than a manual replacement
# table for punctuation.
_CORE_FONTS_ENCODING = "cp1252"


def _sanitize_for_pdf(text: str) -> str:
    # NFKD decomposes ligatures like "\ufb01" (fi) into "f" + "i".
    normalized = unicodedata.normalize("NFKD", text)
    return normalized.encode(_CORE_FONTS_ENCODING, errors="ignore").decode(_CORE_FONTS_ENCODING)


def _sanitize_optional(text: str | None) -> str | None:
    return _sanitize_for_pdf(text) if text else text


def _sanitize_profile_for_pdf(profile: ResumeProfile) -> ResumeProfile:
    return ResumeProfile(
        contact=ContactInfo(
            name=_sanitize_optional(profile.contact.name),
            email=_sanitize_optional(profile.contact.email),
            phone=_sanitize_optional(profile.contact.phone),
            location=_sanitize_optional(profile.contact.location),
        ),
        summary=_sanitize_optional(profile.summary),
        skills=[_sanitize_for_pdf(s) for s in profile.skills],
        experiences=[
            ExperienceEntry(
                company=_sanitize_optional(exp.company),
                title=_sanitize_optional(exp.title),
                start_date=_sanitize_optional(exp.start_date),
                end_date=_sanitize_optional(exp.end_date),
                location=_sanitize_optional(exp.location),
                bullets=[_sanitize_for_pdf(b) for b in exp.bullets],
                detected_skills=exp.detected_skills,
                technologies=exp.technologies,
            )
            for exp in profile.experiences
        ],
        education=[_sanitize_for_pdf(e) for e in profile.education],
        certifications=[_sanitize_for_pdf(c) for c in profile.certifications],
    )


def render_pdf(profile: ResumeProfile) -> bytes:
    profile = _sanitize_profile_for_pdf(profile)
    pdf = FPDF()
    pdf.core_fonts_encoding = _CORE_FONTS_ENCODING
    pdf.set_margins(left=18, top=16, right=18)
    pdf.set_auto_page_break(auto=True, margin=16)
    pdf.add_page()

    _render_header(pdf, profile)

    if profile.summary:
        _section_heading(pdf, "Summary")
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(*_TEXT_COLOR)
        pdf.multi_cell(0, 5.4, profile.summary, align="L", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

    if profile.skills:
        _section_heading(pdf, "Skills")
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(*_TEXT_COLOR)
        pdf.multi_cell(
            0, 5.4, "   \u2022   ".join(profile.skills), align="L", new_x="LMARGIN", new_y="NEXT"
        )
        pdf.ln(3)

    if profile.experiences:
        _section_heading(pdf, "Experience")
        for i, exp in enumerate(profile.experiences):
            if i > 0:
                pdf.ln(3)

            title_line = " \u00b7 ".join(v for v in [exp.title, exp.company] if v) or "Experience"
            date_line = " - ".join(v for v in [exp.start_date, exp.end_date] if v)
            _title_with_trailing_date(pdf, title_line, date_line)

            if exp.location:
                pdf.set_font("Helvetica", "I", 9)
                pdf.set_text_color(*_MUTED_COLOR)
                pdf.cell(0, 5, exp.location, new_x="LMARGIN", new_y="NEXT")

            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*_TEXT_COLOR)
            for bullet in exp.bullets:
                _bullet_item(pdf, bullet)
        pdf.ln(3)

    if profile.education:
        _section_heading(pdf, "Education")
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(*_TEXT_COLOR)
        for entry in profile.education:
            _bullet_item(pdf, entry)
        pdf.ln(3)

    if profile.certifications:
        _section_heading(pdf, "Certifications")
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(*_TEXT_COLOR)
        for entry in profile.certifications:
            _bullet_item(pdf, entry)

    return bytes(pdf.output())


_ACCENT_COLOR = (31, 61, 92)  # dark navy, used for the name and section headings
_BAND_COLOR = (234, 239, 245)  # light navy tint behind the header block
_TEXT_COLOR = (30, 30, 30)
_MUTED_COLOR = (100, 100, 100)
_RULE_COLOR = (200, 205, 212)


def _render_header(pdf: FPDF, profile: ResumeProfile) -> None:
    content_width = pdf.w - pdf.l_margin - pdf.r_margin
    contact_line = " \u00b7 ".join(
        v for v in [profile.contact.email, profile.contact.phone, profile.contact.location] if v
    )
    band_height = 28 if contact_line else 20

    pdf.set_fill_color(*_BAND_COLOR)
    pdf.rect(0, 0, pdf.w, band_height, style="F")

    pdf.set_xy(pdf.l_margin, 8)
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(*_ACCENT_COLOR)
    name_text = profile.contact.name or "Resume"
    pdf.cell(content_width, 10, name_text, align="C", new_x="LMARGIN", new_y="NEXT")

    if contact_line:
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(*_MUTED_COLOR)
        pdf.cell(content_width, 6, contact_line, align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.set_xy(pdf.l_margin, band_height + 8)


def _section_heading(pdf: FPDF, text: str) -> None:
    pdf.set_font("Helvetica", "B", 11.5)
    pdf.set_text_color(*_ACCENT_COLOR)
    pdf.cell(0, 6, text.upper(), new_x="LMARGIN", new_y="NEXT")
    pdf.set_draw_color(*_RULE_COLOR)
    pdf.set_line_width(0.25)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(2.5)


def _title_with_trailing_date(pdf: FPDF, title_text: str, date_text: str) -> None:
    """Renders a bold title on the left and an italic date range flush right
    on the same line — the conventional resume experience-entry layout."""
    content_width = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.set_font("Helvetica", "I", 9.5)
    date_width = pdf.get_string_width(date_text) + 2 if date_text else 0
    title_width = content_width - date_width

    pdf.set_font("Helvetica", "B", 10.5)
    pdf.set_text_color(*_TEXT_COLOR)
    pdf.cell(title_width, 6, title_text, new_x="RIGHT", new_y="TOP")

    if date_text:
        pdf.set_font("Helvetica", "I", 9.5)
        pdf.set_text_color(*_MUTED_COLOR)
        pdf.cell(date_width, 6, date_text, align="R", new_x="LMARGIN", new_y="NEXT")
    else:
        pdf.ln(6)


def _bullet_item(pdf: FPDF, text: str, indent: float = 5.5) -> None:
    """Renders a bulleted line with a hanging indent, so wrapped lines align
    under the text rather than under the bullet marker."""
    x_start = pdf.l_margin
    pdf.set_x(x_start)
    pdf.cell(indent, 5.4, "\u2022", new_x="RIGHT", new_y="TOP")
    pdf.set_x(x_start + indent)
    available_width = pdf.w - pdf.r_margin - (x_start + indent)
    pdf.multi_cell(available_width, 5.4, text, align="L", new_x="LMARGIN", new_y="NEXT")


def render_docx(profile: ResumeProfile) -> bytes:
    doc = Document()
    _set_docx_base_style(doc)

    name_heading = doc.add_heading(profile.contact.name or "Resume", level=1)
    name_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_line = " \u00b7 ".join(
        v for v in [profile.contact.email, profile.contact.phone, profile.contact.location] if v
    )
    if contact_line:
        contact_para = doc.add_paragraph(contact_line)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    if profile.summary:
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(profile.summary)

    if profile.skills:
        doc.add_heading("Skills", level=2)
        doc.add_paragraph("  \u00b7  ".join(profile.skills))

    if profile.experiences:
        doc.add_heading("Experience", level=2)
        for exp in profile.experiences:
            title_line = " \u2014 ".join(v for v in [exp.title, exp.company] if v) or "Experience"
            date_line = " - ".join(v for v in [exp.start_date, exp.end_date] if v)
            _add_docx_title_with_trailing_date(doc, title_line, date_line)

            if exp.location:
                location_para = doc.add_paragraph(exp.location)
                for run in location_para.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    if profile.education:
        doc.add_heading("Education", level=2)
        for entry in profile.education:
            doc.add_paragraph(entry, style="List Bullet")

    if profile.certifications:
        doc.add_heading("Certifications", level=2)
        for entry in profile.certifications:
            doc.add_paragraph(entry, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_docx_base_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for level in (1, 2, 3):
        heading = doc.styles[f"Heading {level}"]
        heading.font.color.rgb = RGBColor(*_ACCENT_COLOR)
        heading.font.name = "Calibri"


def _add_docx_title_with_trailing_date(doc: Document, title_text: str, date_text: str) -> None:
    """Renders a bold title flush left and an italic date range flush right
    on the same line, using a right tab stop at the page's text width."""
    para = doc.add_paragraph()
    section = doc.sections[0]
    text_width = section.page_width - section.left_margin - section.right_margin
    para.paragraph_format.tab_stops.add_tab_stop(text_width, WD_TAB_ALIGNMENT.RIGHT)

    title_run = para.add_run(title_text)
    title_run.bold = True

    if date_text:
        date_run = para.add_run(f"\t{date_text}")
        date_run.italic = True
        date_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
